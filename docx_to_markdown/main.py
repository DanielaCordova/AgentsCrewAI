import os
import tempfile
import uuid
import docling
import litellm
import docx
from crewai import Agent, Task, Crew, Process
from llama_index.readers.docling import DoclingReader

class OllamaLLM:
    """Custom LLM Wrapper for CrewAI using LiteLLM"""
    def __init__(self, model_name="ollama/llama3", api_base="http://localhost:11434"):
        self.model_name = model_name
        self.api_base = api_base

    def __call__(self, prompt: str) -> str:
        """ Calls Ollama through LiteLLM and returns the response in Spanish """
        prompt = f"Responde en español: {prompt}"
        response = litellm.completion(
            model=self.model_name,
            messages=[{"role": "user", "content": prompt}],
            api_base=self.api_base
        )
        return response["choices"][0]["message"]["content"]

# Create Ollama model instance
ollama_llm = OllamaLLM()

def extract_text_from_file(file_path: str) -> str:
    """ Extracts text from either a PDF or DOCX file, preserving structure. """
    if file_path.endswith(".pdf"):
        reader = DoclingReader()
        documents = reader.load_data(file_path)
        return "\n".join([doc.text for doc in documents])
    elif file_path.endswith(".docx"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")

def extract_text_from_docx(file_path: str) -> str:
    """ Extracts structured text from a DOCX file, preserving tables, lists, and headings. """
    doc = docx.Document(file_path)
    output = []
    
    for para in doc.paragraphs:
        style = para.style.name.lower()
        text = para.text.strip()
        
        if "heading" in style:
            level = style.replace("heading ", "")
            output.append(f"{'#' * int(level)} {text}")
        elif para.style.name.startswith("List"):
            output.append(f"- {text}")
        elif text:
            output.append(text)
    
    # Extract Tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append("| " + " | ".join([cell.text.strip() for cell in row.cells]) + " |")
        
        if rows:
            output.append("\n" + rows[0])  # Header row
            output.append("|" + " --- |" * len(table.rows[0].cells))  # Divider
            output.extend(rows[1:])  # Data rows
    
    return "\n".join(output)

# Define document processing agents
doc_extractor = Agent(
    role="Document Extractor",
    goal="Extract text and structure from PDF or DOCX files in Spanish, preserving headings, tables, and lists.",
    backstory="You specialize in parsing documents and retrieving structured content in Spanish, maintaining the original formatting. DO NOT ADD ANY INFORMATION. MANTAIN THE <TODO>",
    verbose=True,
    allow_delegation=False,
    llm=ollama_llm
)

markdown_converter = Agent(
    role="Markdown Converter",
    goal="Convert extracted text into a well-structured Markdown document with properly formatted tables and headers, ensuring correct spacing and proper table formatting without enclosing everything in code blocks. All text must be in Spanish.",
    backstory="You are skilled at formatting documents into clean Markdown with enhanced readability, ensuring proper alignment of table elements, and processing everything in Spanish.",
    verbose=True,
    allow_delegation=False,
    llm=ollama_llm
)

# Define tasks
def extract_doc_task(file_path):
    text = extract_text_from_file(file_path)  # Extract text before defining the task
    return Task(
        description=f"Extract structured text from the document file: {file_path} (in Spanish)",
        agent=doc_extractor,
        expected_output=text,  # Pass extracted text
        config={"file_path": file_path}
    )

def convert_to_markdown_task(extracted_text, output_path="output.md"):
    return Task(
        description="Convert the extracted structured text into a properly formatted Markdown document with structured tables and properly aligned headers.",
        agent=markdown_converter,
        expected_output=f"Structured Markdown file saved at {output_path}, ensuring proper table formatting in Spanish.",
        config={"extracted_text": extracted_text, "output_file": output_path}
    )
# File path simulation (Replace with actual file input)
doc_file_path = "docx_to_markdown/docuHomo.docx" # Can be either a .pdf or .docx file

# Create CrewAI workflow
extracted_text = extract_text_from_file(doc_file_path)  # Extract before CrewAI task execution
extract_task = extract_doc_task(doc_file_path)
convert_task = convert_to_markdown_task(extracted_text)

crew = Crew(
    agents=[doc_extractor, markdown_converter],
    tasks=[extract_task, convert_task],
    verbose=True,
    process=Process.sequential
)

# Execute workflow
output = crew.kickoff()
print(output)


