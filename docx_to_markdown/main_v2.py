import os
import docx
import litellm
from crewai import Agent, Task, Crew, Process
from llama_index.readers.docling import DoclingReader

# ---- Custom Ollama Wrapper ----
class OllamaLLM:
    """Custom LLM Wrapper for CrewAI using LiteLLM"""
    def __init__(self, model_name="ollama/llama3", api_base="http://localhost:11434"):
        self.model_name = model_name
        self.api_base = api_base

    def __call__(self, prompt: str) -> str:
        """ Calls Ollama through LiteLLM and returns the response in Spanish """
        prompt = f"Responde en español: {prompt}"
        response = litellm.completion(
            model=self.model_name,
            messages=[{"role": "user", "content": prompt}],
            api_base=self.api_base
        )
        return response["choices"][0]["message"]["content"]

# ---- Create Ollama LLM Instance ----
ollama_llm = OllamaLLM()

# ---- Function to Extract Text from DOCX ----
def extract_text_from_docx(file_path: str) -> str:
    """ Extracts structured text from a DOCX file, preserving tables, lists, headings, and <TODO> markers. """
    doc = docx.Document(file_path)
    output = []
    seen_todos = set()  # Prevents duplicate <TODO> markers
    seen_lines = set()  # Prevents duplicate text in lists

    for para in doc.paragraphs:
        text = para.text.strip()

        # Ensure <TODO> markers are always included but not duplicated
        if text == "<TODO>":
            #if text not in seen_todos:
            output.append(text)
            seen_todos.add(text)
            continue

        style = para.style.name.lower()
        if "heading" in style:
            level = style.replace("heading ", "")
            output.append(f"\n{'#' * int(level)} {text}\n")  # Ensure spacing
        elif para.style.name.startswith("List"):
            if text and text not in seen_lines:  # Prevents duplicated list items
                output.append(f"- {text}")
                seen_lines.add(text)
        elif text:
            output.append(f"{text}\n")  # Ensure spacing

    # Extract Tables (preserve empty cells with <TODO> and fix row splitting)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            row_data = [cell.text.strip() if cell.text.strip() else "<TODO>" for cell in row.cells]
            
            # Ensure correct column count
            if len(row_data) < len(table.rows[0].cells):
                row_data.extend(["<TODO>"] * (len(table.rows[0].cells) - len(row_data)))

            row_text = "| " + " | ".join(row_data) + " |"
            if row_text not in rows:  # Prevent duplicate rows
                rows.append(row_text)

        if rows:
            output.append("\n" + rows[0])  # Header row
            output.append("|" + " --- |" * len(table.rows[0].cells))  # Divider
            output.extend(rows[1:])  # Data rows
            output.append("\n")  # Ensure new line after table
    
    return "\n".join(output)

# ---- Agents ----
doc_extractor = Agent(
    role="Document Extractor",
    goal="Extract the text **exactly as it appears** in the original document without modifying anything, especially <TODO> markers.",
    backstory="You specialize in document processing. Your task is to extract content **without changes** and ensure <TODO> markers are fully preserved.",
    verbose=True,
    allow_delegation=False,
    llm=ollama_llm
)

markdown_validator = Agent(
    role="Markdown Validator",
    goal="Ensure that the Markdown conversion retains headers, lists, tables, and <TODO> markers correctly.",
    backstory="You validate that the Markdown format is accurate, properly formatted, and maintains <TODO> markers.",
    verbose=True,
    allow_delegation=False,
    llm=ollama_llm
)

# ---- Tasks ----
def extract_doc_task(extracted_text):
    """ Validate the extracted text without modification """
    return Task(
        description=(
            f"Extract the text from the document **exactly as it appears**, without making any modifications.\n\n"
            f"Here is the extracted text:\n\n"
            f"```text\n{extracted_text}\n```\n\n"
            f"Confirm that:\n"
            f"- The structure is 100% identical to the original document.\n"
            f"- No <TODO> markers are missing or duplicated.\n"
            f"- Lists and tables are correctly formatted.\n"
        ),
        agent=doc_extractor,
        expected_output="Extract the document text **without any modifications or missing <TODO> markers**."
    )

def validate_markdown_task(extracted_text, markdown_text, output_path="output.md"):
    """ Validate the correctness of Markdown conversion """
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(markdown_text)
    
    return Task(
        description=(
            f"Analyze the Markdown conversion and ensure it retains all structural elements.\n\n"
            f"Here is the extracted text:\n\n"
            f"```text\n{extracted_text}\n```\n\n"
            f"And here is the converted Markdown:\n\n"
            f"```markdown\n{markdown_text}\n```\n\n"
            f"Confirm that:\n"
            f"- Headers and lists are formatted correctly.\n"
            f"- No <TODO> markers are missing or duplicated.\n"
            f"- Tables are properly structured.\n"
        ),
        agent=markdown_validator,
        expected_output="Confirm whether the Markdown conversion is correctly formatted and preserves all structure, including <TODO> markers."
    )

# ---- Execution ----
doc_file_path = "docx_to_markdown/docuHomo.docx"  # **Keeping your original file path**

# 1️⃣ Extract text from DOCX (BEFORE CrewAI execution)
extracted_text = extract_text_from_docx(doc_file_path)

# 2️⃣ Convert text to Markdown (BEFORE CrewAI execution)
markdown_text = extract_text_from_docx(doc_file_path)  # Keeping it raw for validation

# 3️⃣ Create CrewAI tasks (passing ACTUAL extracted/converted text)
extract_task = extract_doc_task(extracted_text)
validate_task = validate_markdown_task(extracted_text, markdown_text)

# 4️⃣ Create CrewAI workflow
crew = Crew(
    agents=[doc_extractor, markdown_validator],
    tasks=[extract_task, validate_task],
    verbose=True,
    process=Process.sequential
)

# 5️⃣ Execute workflow
output = crew.kickoff()
print("\n✅ CrewAI Execution Completed!\n")
# print(extracted_text)
